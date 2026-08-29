# -*- coding: utf-8 -*-
"""Genera Guia_Curso_Editable_Interactivo.docx"""
import os
import re
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "content",
    "excel",
    "Guia_Curso_Editable_Interactivo.docx",
)
XLSX = os.path.join(os.path.dirname(OUTPUT), "Curso_Editable_Interactivo.xlsx")


def extract_sheet_map(xlsx_path: str) -> list[tuple[str, str, bool]]:
    """Nombre hoja, archivo xml, oculta."""
    rows = []
    with zipfile.ZipFile(xlsx_path) as z:
        wb = z.read("xl/workbook.xml").decode("utf-8", errors="replace")
        rels = z.read("xl/_rels/workbook.xml.rels").decode("utf-8", errors="replace")
        rid_file = {}
        for m in re.finditer(r'Id="(rId\d+)"[^>]*Target="worksheets/(sheet\d+\.xml)"', rels):
            rid_file[m.group(1)] = m.group(2)
        for m in re.finditer(r"<sheet ([^/]+)/>", wb):
            attrs = m.group(1)
            name_m = re.search(r'name="([^"]*)"', attrs)
            rid_m = re.search(r'r:id="(rId\d+)"', attrs)
            if not name_m or not rid_m:
                continue
            hidden = "state=\"hidden\"" in attrs
            rows.append((name_m.group(1), rid_file.get(rid_m.group(1), "?"), hidden))
    return rows


def add_title(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(11)

    # Portada
    t = doc.add_heading("Guía práctica", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Curso_Editable_Interactivo.xlsx")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = RGBColor(0, 180, 216)
    doc.add_paragraph(
        "Cómo funciona el menú interactivo, cómo agregar o quitar pestañas, "
        "y dónde están los elementos «ocultos» del archivo."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("veloX · Sección Excel · Minicursos").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Qué es este archivo
    add_title(doc, "1. ¿Qué es este archivo y por qué es diferente?")
    doc.add_paragraph(
        "El archivo Curso_Editable_Interactivo.xlsx es una copia editable del curso "
        "original (Curso.xlsx) en la que se eliminó la protección con contraseña de las hojas, "
        "pero se conservaron todos los botones del menú superior (234 hipervínculos en formas)."
    )
    doc.add_paragraph(
        "IMPORTANTE: La navegación tipo «aplicación» NO usa las pestañas inferiores de Excel "
        "como menú principal. Usa rectángulos (formas) en las filas 1 y 2 de cada hoja, "
        "con hipervínculos internos que saltan a otra hoja al hacer clic."
    )
    add_bullets(
        doc,
        [
            "Pestañas inferiores = hojas de cálculo (Area, Comp, CadFun, etc.).",
            "Menú negro/azul superior = formas con hipervínculo (#Hoja!A1).",
            "Hojas ocultas DCargos y Auxiliar = datos de soporte (no aparecen en el menú).",
        ],
    )

    # 2. Arquitectura
    add_title(doc, "2. Arquitectura del archivo (cómo funciona por dentro)")
    add_title(doc, "2.1. Tres capas que debes conocer", level=2)
    add_numbered(
        doc,
        [
            "Hojas de cálculo (worksheets): cada pestaña inferior es una hoja con contenido "
            "(texto, tablas, gráficos).",
            "Capa visual del menú (drawings): en cada hoja hay un archivo de dibujo (drawing) "
            "con rectángulos. Cada rectángulo tiene un hipervínculo interno, por ejemplo #Area!A1.",
            "Datos ocultos: hojas DCargos y Auxiliar están ocultas; pueden contener listas, "
            "códigos de cargos o tablas auxiliares usadas por fórmulas.",
        ],
    )
    add_title(doc, "2.2. Flujo al hacer clic en un botón del menú", level=2)
    add_numbered(
        doc,
        [
            "Haces clic en un rectángulo del menú (ej. «2 CONTROL DE FUNCIONARIOS»).",
            "Excel ejecuta el hipervínculo interno (ej. #CFun!A1).",
            "Excel activa la hoja CFun y posiciona la vista en la celda A1.",
            "La hoja destino muestra su propio menú (las mismas formas copiadas en cada hoja).",
        ],
    )
    doc.add_paragraph(
        "Por eso, al navegar, siempre ves el menú superior: cada hoja «importante» "
        "tiene su copia del menú con los mismos botones."
    )

    # 3. Mapa de hojas
    add_title(doc, "3. Mapa de hojas del archivo")
    if os.path.exists(XLSX):
        sheet_rows = extract_sheet_map(XLSX)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Nº"
        hdr[1].text = "Nombre pestaña"
        hdr[2].text = "Archivo interno"
        hdr[3].text = "Visible"
        for i, (name, file, hidden) in enumerate(sheet_rows, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = name
            row[2].text = file
            row[3].text = "Oculta" if hidden else "Visible"
    else:
        doc.add_paragraph("(Abre el archivo xlsx en la misma carpeta para ver el mapa actualizado.)")

    doc.add_paragraph()
    add_title(doc, "3.1. Correspondencia aproximada menú → hoja", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Botón del menú (ejemplo)"
    table2.rows[0].cells[1].text = "Hoja destino"
    menu_map = [
        ("1. INICIO / 1.1 ENFOQUE", "Area"),
        ("1.2 ESQUEMA ORG. / comparativos", "Comp, CritDec, Anual"),
        ("2 CONTROL DE FUNCIONARIOS", "CFun, CadFun, Faltas, Atrasos"),
        ("3 CONTROL DE FORMACIÓN", "Trei, Pres, PDes"),
        ("4 EVALUACIÓN - DESEMPEÑO", "Ava, Dias, Pro"),
        ("5 CONSULTA DE FUNCIONARIOS", "Varias / consultas"),
        ("6 INFORMES", "CadFun, Anual, Mensal"),
        ("7 DASHBOARDS", "DGestão, gráficos en varias hojas"),
        ("Datos auxiliares (ocultos)", "DCargos, Auxiliar"),
    ]
    for a, b in menu_map:
        r = table2.add_row().cells
        r[0].text = a
        r[1].text = b

    doc.add_page_break()

    # 4. Abrir y usar
    add_title(doc, "4. Paso a paso: abrir y usar el archivo")
    add_numbered(
        doc,
        [
            "Cierra cualquier copia antigua Curso_Editable.xlsx (versión sin botones).",
            "Abre Curso_Editable_Interactivo.xlsx con Excel de escritorio (Windows o Mac).",
            "Si aparece barra amarilla «Vista protegida», clic en Habilitar edición.",
            "Ve a la hoja Area (pestaña inferior) o la que tengas abierta.",
            "Haz clic en los rectángulos del menú (filas 1–2), no en celdas vacías.",
            "Para editar contenido: haz clic en celdas debajo del menú (MISIÓN, VISIÓN, etc.).",
        ],
    )

    # 5. Agregar pestaña
    add_title(doc, "5. Paso a paso: AGREGAR una nueva pestaña (hoja + menú)")
    add_title(doc, "5.1. Crear la hoja nueva", level=2)
    add_numbered(
        doc,
        [
            "Clic derecho en cualquier pestaña inferior → Insertar → Hoja de cálculo.",
            "Clic derecho en la nueva pestaña → Cambiar nombre (ej. MiModulo).",
            "Usa nombres cortos sin espacios si vas a enlazar mucho (ej. ModRRHH).",
            "Diseña el contenido desde la fila 5 o inferior (deja filas 1–4 para el menú).",
        ],
    )
    add_title(doc, "5.2. Copiar el menú interactivo a la nueva hoja", level=2)
    add_numbered(
        doc,
        [
            "Ve a una hoja que ya tenga el menú (ej. Area).",
            "Mantén Ctrl y haz clic en cada rectángulo del menú (filas 1–2) para seleccionar "
            "todas las formas, o selecciónalas una por una.",
            "Alternativa: clic en una forma → Ctrl+A si solo hay formas en la hoja.",
            "Ctrl+C para copiar.",
            "Ve a tu hoja nueva → clic en celda A1 → Ctrl+V.",
            "Ajusta posición si hace falta (las formas quedan en la misma zona).",
        ],
    )
    add_title(doc, "5.3. Agregar un botón nuevo que apunte a tu hoja", level=2)
    add_numbered(
        doc,
        [
            "En la hoja Area (menú principal), selecciona un rectángulo similar al que quieres.",
            "Ctrl+C y Ctrl+V para duplicar el rectángulo.",
            "Arrastra el nuevo rectángulo a un espacio libre en la fila 1 o 2.",
            "Clic derecho en la forma → Editar texto → escribe el nombre (ej. «8 MI MÓDULO»).",
            "Clic derecho en la forma → Editar hipervínculo (o Ctrl+K).",
            "Selecciona «Lugar de este documento».",
            "En «Referencia de celda», elige tu hoja (MiModulo) y celda $A$1.",
            "Aceptar.",
        ],
    )
    add_title(doc, "5.4. Repetir el enlace en TODAS las hojas con menú", level=2)
    doc.add_paragraph(
        "Para que el nuevo botón funcione desde cualquier pantalla, debes copiar ese "
        "rectángulo nuevo (o actualizar hipervínculos) en el menú de cada hoja que ya "
        "tenía menú, o copiar de nuevo el menú completo desde Area a las demás hojas."
    )
    add_bullets(
        doc,
        [
            "Método rápido: copia solo el botón nuevo y pégalo en fila 1–2 de cada hoja.",
            "Método limpio: rediseña el menú en Area y vuelve a copiar todo el bloque de formas "
            "a las otras hojas (más trabajo, más consistente).",
        ],
    )

    # 6. Eliminar pestaña
    add_title(doc, "6. Paso a paso: ELIMINAR una pestaña o botón del menú")
    add_title(doc, "6.1. Eliminar solo un botón del menú (sin borrar la hoja)", level=2)
    add_numbered(
        doc,
        [
            "Ve a la hoja donde está el botón.",
            "Clic en el rectángulo del menú que quieres quitar.",
            "Pulsa Supr o Delete.",
            "Repite en cada hoja que tenga copia de ese botón.",
        ],
    )
    add_title(doc, "6.2. Eliminar una hoja completa", level=2)
    add_numbered(
        doc,
        [
            "Primero elimina o redirige los hipervínculos que apuntan a esa hoja "
            "(botones del menú en otras hojas). Si no, los clics darán error.",
            "Clic derecho en la pestaña inferior de la hoja → Eliminar.",
            "Confirma. Excel no permite eliminar la última hoja del libro.",
        ],
    )
    add_title(doc, "6.3. Ocultar una hoja sin eliminarla (como DCargos)", level=2)
    add_numbered(
        doc,
        [
            "Clic derecho en la pestaña → Ocultar.",
            "Para verla de nuevo: Revisar → Mostrar → Ocultar/Mostrar hojas → marcar la hoja.",
            "Útil para datos auxiliares que no deben verse en el menú.",
        ],
    )

    doc.add_page_break()

    # 7. Códigos ocultos
    add_title(doc, "7. Dónde están los «códigos ocultos» y cómo verlos")
    doc.add_paragraph(
        "No hay macros VBA en este archivo. Lo «oculto» son hojas, hipervínculos en formas, "
        "fórmulas y archivos internos del .xlsx."
    )
    add_title(doc, "7.1. Hojas ocultas", level=2)
    add_numbered(
        doc,
        [
            "Revisar → Mostrar → Ocultar/Mostrar hojas.",
            "En este curso: DCargos y Auxiliar suelen estar ocultas.",
            "Ahí pueden estar tablas de cargos, listas para validación o códigos internos.",
        ],
    )
    add_title(doc, "7.2. Hipervínculos de los botones (el «código» de navegación)", level=2)
    add_numbered(
        doc,
        [
            "Selecciona un rectángulo del menú.",
            "Clic derecho → Editar hipervínculo.",
            "Verás algo como: Lugar en este documento → CFun → Celda A1.",
            "Eso equivale al código interno #CFun!A1.",
            "Para listar todos: en una hoja, repite con cada botón, o usa Desarrollador → "
            "Visual Basic solo si en el futuro agregas macros.",
        ],
    )
    add_title(doc, "7.3. Ver todas las formas del menú", level=2)
    add_numbered(
        doc,
        [
            "Inicio → Buscar y seleccionar → Panel de selección (o Formato → Panel de selección).",
            "Se abre una lista de todas las formas en la hoja activa (Rectángulo 1, 2, 3…).",
            "Ahí puedes seleccionar, ocultar o eliminar formas sin buscarlas con el mouse.",
        ],
    )
    add_title(doc, "7.4. Fórmulas y referencias a hojas ocultas", level=2)
    add_numbered(
        doc,
        [
            "En cualquier celda con fórmula, clic en la celda y revisa la barra de fórmulas.",
            "Referencias como =Auxiliar!A1 o =DCargos!B5 apuntan a hojas ocultas.",
            "Buscar en hoja (Ctrl+F) → pestaña «Opciones» → Buscar en: Hoja o Libro.",
            "Busca textos como DCargos, Auxiliar, INDIRECTO, BUSCARV, BUSCARX.",
        ],
    )
    add_title(doc, "7.5. Archivos internos del .xlsx (nivel avanzado)", level=2)
    doc.add_paragraph(
        "Un .xlsx es una carpeta comprimida. Si cambias el nombre a .zip y abres, verás:"
    )
    add_bullets(
        doc,
        [
            "xl/workbook.xml → lista de hojas y cuáles están ocultas.",
            "xl/worksheets/sheetN.xml → contenido de cada hoja.",
            "xl/drawings/drawingN.xml → formas del menú (rectángulos).",
            "xl/drawings/_rels/drawingN.xml.rels → destinos #Hoja!A1 de cada botón.",
            "xl/sharedStrings.xml → textos repetidos (MISIÓN, VISIÓN, etc.).",
        ],
    )
    doc.add_paragraph(
        "No edites estos archivos manualmente salvo que sepas XML; usa Excel para cambios normales."
    )

    # 8. Buenas prácticas
    add_title(doc, "8. Buenas prácticas y errores comunes")
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Table Grid"
    table3.rows[0].cells[0].text = "Error común"
    table3.rows[0].cells[1].text = "Solución"
    errors = [
        ("No puedo clicar el menú", "Usa Curso_Editable_Interactivo.xlsx, no la versión vieja "
         "Curso_Editable.xlsx. Habilitar edición si hay Vista protegida."),
        ("Clic en botón no lleva a ningún lado", "Editar hipervínculo de la forma; verificar "
         "que la hoja destino existe y el nombre coincide."),
        ("Al guardar perdí los botones", "No guardes con herramientas que reescriben el xlsx "
         "(algunos conversores online). Guarda con Excel."),
        ("No puedo editar celdas", "Revisar → Desproteger hoja (en Interactivo ya está desprotegido)."),
        ("Quiero volver a generar el editable", "Ejecutar: python scripts/unlock_curso_xlsx.py "
         "desde la carpeta del proyecto."),
    ]
    for e, s in errors:
        r = table3.add_row().cells
        r[0].text = e
        r[1].text = s

    # 9. Resumen
    add_title(doc, "9. Resumen rápido")
    add_bullets(
        doc,
        [
            "Menú superior = formas con hipervínculo, no pestañas de Excel.",
            "Agregar pestaña = Insertar hoja + copiar menú + Editar hipervínculo del botón nuevo.",
            "Eliminar pestaña = quitar hipervínculos primero, luego Eliminar hoja.",
            "Ocultos = hojas DCargos/Auxiliar + hipervínculos en formas + fórmulas.",
            "Panel de selección = ver todas las formas de una hoja.",
            "Archivo a usar siempre: content/excel/Curso_Editable_Interactivo.xlsx",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Documento generado para veloX · Distribuidora Andina / Curso de gestión")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generado: {OUTPUT}")


if __name__ == "__main__":
    build_document()
